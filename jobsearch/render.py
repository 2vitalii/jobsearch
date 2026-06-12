"""Résumé rendering and application-kit assembly.

Everything here returns in-memory artifacts (``bytes`` / ``str``) — nothing is
written to disk. Where the kit lands (a local ``review/`` folder now, object
storage later) is the orchestration/store layer's decision.

``safe_name`` is a security control, not a convenience: every scraped/untrusted
string that becomes part of a filesystem path goes through it (path-traversal
defence). Keep it strict: ``[\\w\\-]`` only, length-capped.
"""

from __future__ import annotations

import re
import sys
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Нет python-docx. Установи: pip install python-docx")

from .models import Job, MatchResult, Package

BLUE = RGBColor(0x2E, 0x5A, 0x8C)
GRAY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Мастер-CV: парсинг строки в структуру (CV приходит текстом, диск не читаем)
# ---------------------------------------------------------------------------
def parse_master(txt: str):
    lines = txt.splitlines()
    name, extras, sections, cur = "", [], [], None
    i = 0
    while i < len(lines):
        if lines[i].startswith("# ") and not lines[i].startswith("## "):
            name = lines[i][2:].strip()
            i += 1
            break
        i += 1
    while i < len(lines) and not lines[i].startswith("## "):
        if lines[i].strip():
            extras.append(lines[i].strip())
        i += 1
    while i < len(lines):
        l = lines[i]
        if l.startswith("## "):
            cur = {"title": l[3:].strip(), "items": []}
            sections.append(cur)
            i += 1
            continue
        if cur is None:
            i += 1
            continue
        if l.startswith("### "):
            cur["items"].append(("h3", l[4:].strip()))
            i += 1
            j = i
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and not lines[j].startswith(("- ", "### ", "## ")):
                cur["items"].append(("meta", lines[j].strip()))
                i = j + 1
            continue
        if l.strip().startswith("- "):
            cur["items"].append(("bullet", l.strip()[2:].strip()))
            i += 1
            continue
        if l.strip():
            cur["items"].append(("para", l.strip()))
            i += 1
            continue
        i += 1
    return name, extras, sections


def apply_tailoring(sections, summary: str, skills: list):
    for s in sections:
        t = s["title"].lower()
        if t.startswith("professional summary") and summary.strip():
            s["items"] = [("para", summary.strip())]
        elif t.startswith("core skills") and skills:
            s["items"] = [("bullet", x.strip().lstrip("-").strip()) for x in skills if x.strip()]
    return sections


# ---------------------------------------------------------------------------
# Рендер резюме в .docx (-> bytes)
# ---------------------------------------------------------------------------
def _bottom_border(p, color="2E5A8C", size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    # pBdr должен идти ДО spacing/ind/jc/rPr по схеме CT_PPr
    ref = None
    for tag in ("w:spacing", "w:ind", "w:jc", "w:rPr"):
        el = pPr.find(qn(tag))
        if el is not None:
            ref = el
            break
    if ref is not None:
        ref.addprevious(pbdr)
    else:
        pPr.append(pbdr)


def render_cv(name, extras, sections) -> bytes:
    """Render the tailored résumé to a .docx and return the raw bytes."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.75)
        sec.left_margin = sec.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(2)

    if extras:
        p = doc.add_paragraph()
        r = p.add_run(extras[0])
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        p.paragraph_format.space_after = Pt(2)
        for ex in extras[1:]:
            p = doc.add_paragraph()
            r = p.add_run(ex)
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

    for s in sections:
        h = doc.add_paragraph()
        r = h.add_run(s["title"].upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(5)
        _bottom_border(h)
        bold_leads = s["title"].strip().lower() in ("core skills", "additional information")
        for kind, text in s["items"]:
            if kind == "h3":
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
            elif kind == "meta":
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = GRAY
                p.paragraph_format.space_after = Pt(3)
            elif kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                if bold_leads and ": " in text:
                    lead, rest = text.split(": ", 1)
                    rr = p.add_run(lead + ": ")
                    rr.bold = True
                    rr.font.size = Pt(10.5)
                    rr2 = p.add_run(rest)
                    rr2.font.size = Pt(10.5)
                else:
                    rr = p.add_run(text)
                    rr.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(4)
    # python-docx-шаблон оставляет <w:zoom> без percent — добавляем, чтобы файл был валидным
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Отчёт + сборка комплекта
# ---------------------------------------------------------------------------
def safe_name(s: str) -> str:
    return re.sub(r"[^\w\-]+", "_", s)[:50].strip("_")


def report_md(job: Job, res: MatchResult) -> str:
    def lst(x):
        return ", ".join(x) if isinstance(x, list) else str(x)
    return (
        f"# {job.title} — {job.company}\n\n"
        f"- Fit: **{res.fit_score}/100** | B2B: **{res.b2b}**\n"
        f"- {res.reason}\n"
        f"- Ссылка: {job.url}\n\n"
        f"## Вердикт рекрутёра\n{res.recruiter_verdict}\n\n"
        f"## ATS-match\n"
        f"**Ключевые слова вакансии:** {lst(res.jd_keywords)}\n\n"
        f"**Есть у тебя (поднято в резюме):** {lst(res.ats_present)}\n\n"
        f"**Не хватает:** {lst(res.ats_missing)}\n\n"
        f"## Пробелы (что переформулировать / доучить)\n{res.gaps}\n"
    )


def build_package(job: Job, res: MatchResult, cv_text: str) -> Package:
    """Assemble the full application kit in memory: tailored CV (docx bytes),
    cover letter, ATS report. Writing to disk is the orchestration's job."""
    name, extras, sections = parse_master(cv_text)
    apply_tailoring(sections, res.tailored_summary, res.tailored_skills or [])
    cv_docx = render_cv(name, extras, sections)
    cover = res.cover_letter.strip() + "\n"
    report = report_md(job, res)
    return Package(cv_docx=cv_docx, cover_letter=cover, ats_report=report)
