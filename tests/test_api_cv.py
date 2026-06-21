"""Integration tests for CV intake (/cv/upload, GET /cv, PUT /cv).

SKIPPED by default, same gate as the other Supabase integration tests: needs the
``supabase`` + ``fastapi`` packages (and ``python-docx`` for building the upload)
AND a live project (SUPABASE_URL + SUPABASE_SECRET_KEY in env).

The LLM is NOT called for real: ``get_llm`` is overridden with a deterministic
fake that returns canonical markdown / profile, so the tests are free and stable.
A throwaway auth user is minted on dedicated clients (see test_api_auth) and
deleted in teardown — the cascade removes the cvs row.
"""

from __future__ import annotations

import io
import os
import uuid

import pytest

pytest.importorskip("supabase")
pytest.importorskip("fastapi")
pytest.importorskip("docx")

if not (os.environ.get("SUPABASE_URL") and os.environ.get("SUPABASE_SECRET_KEY")):
    pytest.skip(
        "needs SUPABASE_URL + SUPABASE_SECRET_KEY (live project)",
        allow_module_level=True,
    )

from fastapi.testclient import TestClient

from api.main import app
from api.deps import get_llm
from jobsearch.supabase_store import make_supabase_client

CANON_MD = "# Jane Doe\nSupport Engineer\nEU (Remote) · jane@example.com\n\n## Professional Summary\nSupport engineer."
CANON_MD_2 = "# Jane Doe\nIntegration Engineer\nEU (Remote) · jane@example.com\n\n## Professional Summary\nIntegration engineer."
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class FakeLLM:
    """Deterministic LLM: parse_cv asks with PARSE_SYSTEM, make_short_profile with
    PROFILE_SYSTEM. We branch on which system prompt arrives so each helper gets a
    sensible canned answer — and echo a marker of the input so PUT can be told
    apart from upload."""

    def complete(self, *, model, system, messages, max_tokens) -> str:
        content = messages[0]["content"]
        if "short" in system.lower() or "profile" in system.lower():
            # short profile: echo the headline-ish first words of the CV
            tag = "integration" if "Integration" in content else "support"
            return f"Fake short profile ({tag}), 2 years, B2B."
        # parse_cv: return canonical markdown
        return CANON_MD


@pytest.fixture(scope="module")
def tc():
    app.dependency_overrides[get_llm] = lambda: FakeLLM()
    client = TestClient(app)
    yield client
    app.dependency_overrides.pop(get_llm, None)


@pytest.fixture()
def auth_user():
    admin = make_supabase_client()
    email = f"test_{uuid.uuid4().hex}@example.com"
    password = uuid.uuid4().hex
    created = admin.auth.admin.create_user(
        {"email": email, "password": password, "email_confirm": True}
    )
    uid = created.user.id
    signer = make_supabase_client()
    signed = signer.auth.sign_in_with_password({"email": email, "password": password})
    token = signed.session.access_token
    yield {"user_id": uid, "email": email, "token": token}
    admin.auth.admin.delete_user(uid)  # cascade removes the cvs row


def _tiny_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Support Engineer with 2 years of experience. SQL, REST APIs.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _auth(user) -> dict:
    return {"Authorization": f"Bearer {user['token']}"}


def test_upload_requires_token(tc):
    files = {"file": ("cv.docx", _tiny_docx_bytes(), DOCX_CT)}
    r = tc.post("/cv/upload", files=files)
    assert r.status_code == 401


def test_upload_parses_and_persists(tc, auth_user):
    files = {"file": ("cv.docx", _tiny_docx_bytes(), DOCX_CT)}
    r = tc.post("/cv/upload", files=files, headers=_auth(auth_user))
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["markdown"] == CANON_MD
    assert "Fake short profile" in body["short_profile"]

    # row really landed in cvs
    sb = make_supabase_client()
    row = (
        sb.table("cvs").select("markdown, short_profile")
        .eq("user_id", auth_user["user_id"]).single().execute()
    )
    assert row.data["markdown"] == CANON_MD
    assert "Fake short profile" in row.data["short_profile"]


def test_get_returns_same(tc, auth_user):
    files = {"file": ("cv.docx", _tiny_docx_bytes(), DOCX_CT)}
    tc.post("/cv/upload", files=files, headers=_auth(auth_user))
    r = tc.get("/cv", headers=_auth(auth_user))
    assert r.status_code == 200
    assert r.json()["markdown"] == CANON_MD


def test_get_404_when_absent(tc, auth_user):
    r = tc.get("/cv", headers=_auth(auth_user))
    assert r.status_code == 404


def test_put_regenerates_short_profile(tc, auth_user):
    files = {"file": ("cv.docx", _tiny_docx_bytes(), DOCX_CT)}
    tc.post("/cv/upload", files=files, headers=_auth(auth_user))

    r = tc.put("/cv", json={"markdown": CANON_MD_2}, headers=_auth(auth_user))
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["markdown"] == CANON_MD_2
    # short_profile is regenerated from the NEW markdown (integration, not support)
    assert "integration" in body["short_profile"]
