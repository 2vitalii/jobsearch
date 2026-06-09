#!/usr/bin/env python3
"""
pipeline.py — ATS-движок. По каждой вакансии из jobs_*.csv:
  - читает master_cv.md (источник правды),
  - через Claude API делает разбор вакансии, ATS-match, подгоняет summary и skills,
  - собирает подогнанное резюме (.docx) + сопроводительное + ATS-отчёт,
  - складывает комплект в review/<скор>_<компания>_<вакансия>/.

Тело опыта/проектов/образования НЕ переписывается (защита от вранья) — под вакансию
тюнятся summary и порядок/формулировки skills, где это правда; остальное в ATS-отчёте.

Запуск:
  pip install requests python-docx
  export ANTHROPIC_API_KEY="sk-ant-..."
  python pipeline.py [jobs_YYYY-MM-DD.csv]
"""

import csv
import datetime
import glob
import json
import os
import re
import sys
import time

try:
    import requests
except ImportError:
    sys.exit("Нет requests. Установи: pip install requests python-docx")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Нет python-docx. Установи: pip install python-docx")

# ---------------------------------------------------------------------------
# НАСТРОЙКИ
# ---------------------------------------------------------------------------
API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
MODEL_TAILOR = "claude-sonnet-4-6"            # дорогой шаг: полный тюнинг (можно opus для качества)
MODEL_SCORE = "claude-haiku-4-5-20251001"     # дешёвый шаг: предварительная оценка fit
MAX_JOBS = int(os.environ.get("MAX_JOBS", "250"))  # потолок очереди за прогон (env MAX_JOBS переопределяет)
MIN_FIT = 45                                  # финальный порог (после Sonnet); ниже = больше откликов, слабее средний фит
PRE_MIN_FIT = 35                              # мягкий порог Haiku-гейта; ниже = больше доходит до Sonnet (дороже)
PROCESS_REGIONS = {"WORLDWIDE", "EUROPE", "UNKNOWN"}
REVIEW_DIR = "review"
DONE_LOG = ".processed_urls.txt"
MASTER_CV = "master_cv.md"
APPLICATIONS_CSV = "applications.csv"   # трекер: что сгенерировано / на что откликнулся / ответы

# Короткий профиль для дешёвого предфильтра (Haiku) — не грузим полный CV ради экономии
SHORT_PROFILE = (
    "Technical Support & Integration Engineer, 2 года (Axxonsoft). 70-90 тикетов/нед, "
    ">85% self-resolution. SQL, REST API, MQTT, Azure, Git, Ruby DSL, JasperReports, "
    "Grafana/InfluxDB. Также реальный опыт Project Manager: 2 доведённых проекта, команды "
    "6-8, доставка, стейкхолдеры, кросс-функциональная координация. B2B через ИП в Армении. "
    "EN C1, RU native."
)

BLUE = RGBColor(0x2E, 0x5A, 0x8C)
GRAY = RGBColor(0x55, 0x55, 0x55)

ENDPOINT = "https://api.anthropic.com/v1/messages"
HEADERS = {"x-api-key": API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"}

SYSTEM = (
    "Ты — ATS-эксперт и ассистент по поиску работы. На входе: полный мастер-CV кандидата "
    "и описание вакансии. Оцени соответствие и подготовь материалы для отклика.\n"
    "Кандидат работает по B2B через sole proprietor (IE) в Армении: инвойсы, EOR не нужен. "
    "Если вакансия требует штатного найма в конкретной стране/штате или исключает контракторов "
    "— снижай fit и ставь b2b_eligible no/maybe.\n"
    "\n"
    "ЖЁСТКИЕ ПРАВИЛА ЧЕСТНОСТИ:\n"
    "- У кандидата ДВА реальных трека: (1) Technical Support & Integration Engineer и "
    "(2) реальный опыт Project Manager — 2 доведённых проекта, команды 6-8, работа со "
    "стейкхолдерами и кросс-функциональная координация. Под саппорт/интеграционные вакансии "
    "веди от поддержки; под project/product management — от PM-опыта (управление, доставка, "
    "координация, стейкхолдеры). Под PM-роли ТАКЖЕ подсвечивай PM-смежные элементы саппорт-"
    "опыта в управленческих формулировках: кросс-функциональная координация с разработчиками "
    "и продакт-менеджерами; внедрённое им командное улучшение процесса с эффектом ~50%; "
    "приоритизация конкурирующих задач под SLA; коммуникация со стейкхолдерами и обучение; "
    "владение регулярной отчётностью и автоматизацией. НО ЧЕСТНО: не приписывай формальное "
    "руководство проектами в саппорт-роли и не выдумывай PM-артефакты (бюджеты, roadmap, "
    "discovery, продуктовые метрики), которых не было — недостающее идёт в ats_missing/gaps.\n"
    "- Кандидат — ТЕХНИЧЕСКИЙ специалист. Если вакансия по сути нетехническая (чистый "
    "телефон/чат/колл-центр, ресепшн, продажи) без troubleshooting, SQL/API, интеграций, "
    "скриптов или работы с системами — ставь fit низко и отметь это в reason/gaps; не "
    "натягивай его технический профиль на такую роль.\n"
    "- Используй ТОЛЬКО факты и навыки из мастер-CV. Ничего не выдумывай: ни метрик, ни "
    "технологий, ни опыта. Если вакансия хочет то, чего у кандидата нет — это идёт в "
    "ats_missing и gaps, НИКОГДА в summary/skills/письмо.\n"
    "- tailored_skills: переупорядочь и переформулируй РЕАЛЬНЫЕ навыки так, чтобы наверх "
    "вышли ключевые слова вакансии, которые у кандидата правда есть. Не добавляй новых.\n"
    "- tailored_summary: 3-4 строки под акцент вакансии, с ключевыми словами вакансии, без "
    "стаффинга, только правдивое.\n"
    "- cover_letter: никогда не пиши, что кандидат не подходит. Продавай. Первый абзац — "
    "привязка к компании/миссии. Про ИП ровно одно предложение по шаблону: 'I operate as an "
    "independent contractor through my registered sole proprietorship (IE) in Armenia and am "
    "fully set up for international B2B engagement.' Никакой кириллицы в английском тексте.\n"
    "- Язык summary/skills/cover_letter — язык вакансии (по умолчанию английский).\n"
    "- Формат cover_letter адаптивно: формальная/корпоративная вакансия -> развёрнутое письмо "
    "~250-320 слов с блоком 'Key strengths:' (4-5 буллетов, маппинг навыка на требование, с "
    "инструментами) и подписью имя+телефон+email; стартап/краткое объявление -> короткое до "
    "~150 слов в 3 абзаца, подпись именем. Сомневаешься — короткое.\n"
    "- cover_letter — ЧИСТЫЙ ТЕКСТ без markdown: никаких '**', '#', '`'. Его вставляют в "
    "письмо/форму как есть. Буллеты в формальном варианте начинай символом '• ' (без жирного).\n"
    "\n"
    "Поля reason, gaps, recruiter_verdict — ПО-РУССКИ (это для кандидата). reason начни с "
    "[формальное] или [короткое].\n"
    "\n"
    "Отвечай СТРОГО валидным JSON без markdown. Схема:\n"
    "{"
    '"fit_score": <int 0-100>, '
    '"b2b_eligible": "yes|maybe|no", '
    '"reason": "<1 строка по-русски>", '
    '"jd_keywords": ["<до 15 ключевых слов/требований вакансии>"], '
    '"ats_present": ["<ключевые слова вакансии, которые у кандидата ЕСТЬ>"], '
    '"ats_missing": ["<требуемое, чего у кандидата НЕТ>"], '
    '"tailored_summary": "<текст>", '
    '"tailored_skills": ["<Категория: навыки>", "..."], '
    '"gaps": "<пробелы: что переформулировать, что реально доучить — по-русски>", '
    '"recruiter_verdict": "<shortlist|maybe|reject + почему + быстрые правки, по-русски>", '
    '"cover_letter": "<текст>"'
    "}"
)


# ---------------------------------------------------------------------------
# Память обработанных
# ---------------------------------------------------------------------------
def load_processed() -> set:
    if os.path.exists(DONE_LOG):
        with open(DONE_LOG, encoding="utf-8") as f:
            return {ln.strip() for ln in f if ln.strip()}
    return set()


def mark_processed(url: str):
    with open(DONE_LOG, "a", encoding="utf-8") as f:
        f.write(url + "\n")


# ---------------------------------------------------------------------------
# Мастер-CV: загрузка и парсинг в структуру
# ---------------------------------------------------------------------------
def load_master() -> str:
    if not os.path.exists(MASTER_CV):
        sys.exit(f"Нет {MASTER_CV} рядом со скриптом — положи мастер-CV.")
    with open(MASTER_CV, encoding="utf-8") as f:
        txt = f.read()
    return re.sub(r"<!--.*?-->", "", txt, flags=re.S)  # убираем служебные комментарии


def parse_master(txt: str):
    lines = txt.splitlines()
    name, extras, sections, cur = "", [], [], None
    i = 0
    while i < len(lines):
        if lines[i].startswith("# ") and not lines[i].startswith("## "):
            name = lines[i][2:].strip()
            i += 1
            break
        i += 1
    while i < len(lines) and not lines[i].startswith("## "):
        if lines[i].strip():
            extras.append(lines[i].strip())
        i += 1
    while i < len(lines):
        l = lines[i]
        if l.startswith("## "):
            cur = {"title": l[3:].strip(), "items": []}
            sections.append(cur)
            i += 1
            continue
        if cur is None:
            i += 1
            continue
        if l.startswith("### "):
            cur["items"].append(("h3", l[4:].strip()))
            i += 1
            j = i
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and not lines[j].startswith(("- ", "### ", "## ")):
                cur["items"].append(("meta", lines[j].strip()))
                i = j + 1
            continue
        if l.strip().startswith("- "):
            cur["items"].append(("bullet", l.strip()[2:].strip()))
            i += 1
            continue
        if l.strip():
            cur["items"].append(("para", l.strip()))
            i += 1
            continue
        i += 1
    return name, extras, sections


def apply_tailoring(sections, summary: str, skills: list):
    for s in sections:
        t = s["title"].lower()
        if t.startswith("professional summary") and summary.strip():
            s["items"] = [("para", summary.strip())]
        elif t.startswith("core skills") and skills:
            s["items"] = [("bullet", x.strip().lstrip("-").strip()) for x in skills if x.strip()]
    return sections


# ---------------------------------------------------------------------------
# Рендер резюме в .docx
# ---------------------------------------------------------------------------
def _bottom_border(p, color="2E5A8C", size=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    # pBdr должен идти ДО spacing/ind/jc/rPr по схеме CT_PPr
    ref = None
    for tag in ("w:spacing", "w:ind", "w:jc", "w:rPr"):
        el = pPr.find(qn(tag))
        if el is not None:
            ref = el
            break
    if ref is not None:
        ref.addprevious(pbdr)
    else:
        pPr.append(pbdr)


def render_cv(path, name, extras, sections):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.75)
        sec.left_margin = sec.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(18)
    p.paragraph_format.space_after = Pt(2)

    if extras:
        p = doc.add_paragraph()
        r = p.add_run(extras[0])
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        p.paragraph_format.space_after = Pt(2)
        for ex in extras[1:]:
            p = doc.add_paragraph()
            r = p.add_run(ex)
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

    for s in sections:
        h = doc.add_paragraph()
        r = h.add_run(s["title"].upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(5)
        _bottom_border(h)
        bold_leads = s["title"].strip().lower() in ("core skills", "additional information")
        for kind, text in s["items"]:
            if kind == "h3":
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
            elif kind == "meta":
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.italic = True
                r.font.size = Pt(10)
                r.font.color.rgb = GRAY
                p.paragraph_format.space_after = Pt(3)
            elif kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                if bold_leads and ": " in text:
                    lead, rest = text.split(": ", 1)
                    rr = p.add_run(lead + ": ")
                    rr.bold = True
                    rr.font.size = Pt(10.5)
                    rr2 = p.add_run(rest)
                    rr2.font.size = Pt(10.5)
                else:
                    rr = p.add_run(text)
                    rr.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph()
                r = p.add_run(text)
                r.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(4)
    # python-docx-шаблон оставляет <w:zoom> без percent — добавляем, чтобы файл был валидным
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    doc.save(path)


# ---------------------------------------------------------------------------
# API + отчёт
# ---------------------------------------------------------------------------
def parse_json(text: str) -> dict:
    text = re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.MULTILINE).strip()
    try:
        return json.loads(text, strict=False)   # strict=False — терпим control-символы в строках
    except json.JSONDecodeError:
        start = text.find("{")
        if start == -1:
            raise
        obj, _ = json.JSONDecoder(strict=False).raw_decode(text[start:])
        return obj


def _post(body: dict) -> str:
    r = requests.post(ENDPOINT, headers=HEADERS, json=body, timeout=90)
    r.raise_for_status()
    data = r.json()
    return "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")


def score_fit(job: dict) -> dict:
    """Дешёвый предфильтр на Haiku: только fit_score + b2b + причина. Экономит дорогие вызовы."""
    sys_p = (
        "Ты быстро оцениваешь соответствие кандидата вакансии. Кандидат работает по B2B "
        "через ИП в Армении (контракт, без EOR). Штатный найм в конкретной стране/исключение "
        "контракторов -> снижай fit. Кандидат ТЕХНИЧЕСКИЙ (support/integration + PM): если "
        "вакансия — чисто нетехническая поддержка (телефон/чат/колл-центр, продажи, ресепшн) "
        "без troubleshooting, SQL/API, интеграций, скриптов или работы с системами — это НЕ "
        "его профиль, ставь fit низко. Отвечай ТОЛЬКО JSON: "
        '{"fit_score": <0-100>, "b2b_eligible": "yes|maybe|no", "reason": "<1 строка по-русски>"}'
    )
    user = (
        f"КАНДИДАТ: {SHORT_PROFILE}\n\n"
        f"ВАКАНСИЯ: {job.get('title','')} @ {job.get('company','')} ({job.get('location','')})\n"
        f"{(job.get('description') or '')[:2500]}"
    )
    body = {"model": MODEL_SCORE, "max_tokens": 200, "system": sys_p,
            "messages": [{"role": "user", "content": user}]}
    return parse_json(_post(body))


def analyze(job: dict, master_txt: str) -> dict:
    """Дорогой шаг на Sonnet: полный тюнинг. system+CV вынесены в кешируемые блоки."""
    user_msg = (
        f"ВАКАНСИЯ:\nНазвание: {job.get('title','')}\nКомпания: {job.get('company','')}\n"
        f"Локация: {job.get('location','')}\nСсылка: {job.get('url','')}\n\n"
        f"ОПИСАНИЕ ВАКАНСИИ (используй для ATS-ключевых слов и требований):\n"
        f"{(job.get('description') or '')[:6000]}\n\n"
        "Верни JSON по схеме."
    )
    # system как массив блоков; CV помечен cache_control — статичный префикс кешируется
    # между вакансиями (если суммарно >= минимума токенов модели; иначе просто игнорится).
    system_blocks = [
        {"type": "text", "text": SYSTEM},
        {"type": "text", "text": "МАСТЕР-CV КАНДИДАТА:\n" + master_txt,
         "cache_control": {"type": "ephemeral"}},
    ]
    body = {"model": MODEL_TAILOR, "max_tokens": 4000, "system": system_blocks,
            "messages": [{"role": "user", "content": user_msg}]}
    return parse_json(_post(body))


def safe_name(s: str) -> str:
    return re.sub(r"[^\w\-]+", "_", s)[:50].strip("_")


def report_md(job: dict, res: dict) -> str:
    def lst(x):
        return ", ".join(x) if isinstance(x, list) else str(x)
    return (
        f"# {job.get('title','')} — {job.get('company','')}\n\n"
        f"- Fit: **{res.get('fit_score','?')}/100** | B2B: **{res.get('b2b_eligible','?')}**\n"
        f"- {res.get('reason','')}\n"
        f"- Ссылка: {job.get('url','')}\n\n"
        f"## Вердикт рекрутёра\n{res.get('recruiter_verdict','')}\n\n"
        f"## ATS-match\n"
        f"**Ключевые слова вакансии:** {lst(res.get('jd_keywords',[]))}\n\n"
        f"**Есть у тебя (поднято в резюме):** {lst(res.get('ats_present',[]))}\n\n"
        f"**Не хватает:** {lst(res.get('ats_missing',[]))}\n\n"
        f"## Пробелы (что переформулировать / доучить)\n{res.get('gaps','')}\n"
    )


def log_application(job: dict, res: dict, folder: str):
    """Дописывает строку в applications.csv — воронка: сгенерировано -> отклик -> ответ."""
    new = not os.path.exists(APPLICATIONS_CSV)
    with open(APPLICATIONS_CSV, "a", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        if new:
            w.writerow(["date_generated", "fit", "b2b", "region", "company", "title",
                        "source", "url", "folder", "status", "applied_date", "response"])
        w.writerow([
            datetime.date.today().isoformat(), res.get("fit_score", ""),
            res.get("b2b_eligible", ""), job.get("region", ""), job.get("company", ""),
            job.get("title", ""), job.get("source", ""), job.get("url", ""), folder,
            "GENERATED", "", "",
        ])


def write_package(job: dict, res: dict, master_txt: str) -> str:
    score = int(res.get("fit_score", 0) or 0)
    folder = os.path.join(
        REVIEW_DIR,
        f"{score:03d}_{safe_name(job.get('company','x'))}_{safe_name(job.get('title','x'))}",
    )
    os.makedirs(folder, exist_ok=True)

    name, extras, sections = parse_master(master_txt)
    apply_tailoring(sections, res.get("tailored_summary", ""), res.get("tailored_skills", []) or [])
    render_cv(os.path.join(folder, "Vitalii_Vlasov_CV.docx"), name, extras, sections)

    with open(os.path.join(folder, "cover_letter.txt"), "w", encoding="utf-8") as f:
        f.write(res.get("cover_letter", "").strip() + "\n")
    with open(os.path.join(folder, "ats_report.md"), "w", encoding="utf-8") as f:
        f.write(report_md(job, res))
    return folder


# ---------------------------------------------------------------------------
def main():
    if not API_KEY:
        sys.exit("Нет ANTHROPIC_API_KEY. Задай: export ANTHROPIC_API_KEY=sk-ant-...")

    master_txt = load_master()

    processed = load_processed()

    # argv[2] = папка вывода (иначе review)
    global REVIEW_DIR
    if len(sys.argv) > 2 and sys.argv[2].strip():
        REVIEW_DIR = sys.argv[2].strip()

    # Читаем ВЕСЬ накопленный пул (все jobs_*.csv), а не только свежий файл —
    # чтобы прорабатывать бэклог день за днём и ничего не терять.
    if len(sys.argv) > 1 and sys.argv[1].strip():
        files = [sys.argv[1]]
    else:
        files = sorted(glob.glob("jobs_*.csv"))
        if not files:
            sys.exit("Не найден jobs_*.csv — сначала запусти job_finder.py")

    pool = []
    for fp in files:
        try:
            with open(fp, encoding="utf-8-sig") as f:
                pool += list(csv.DictReader(f))
        except Exception as e:
            print(f"  [пропуск {fp}: {e}]")

    # дедуп пула по ссылке
    seen_urls, jobs = set(), []
    for j in pool:
        u = j.get("url", "")
        if u and u in seen_urls:
            continue
        seen_urls.add(u)
        jobs.append(j)

    # лучшие регионы — вперёд
    order = {"WORLDWIDE": 0, "EUROPE": 1, "UNKNOWN": 2, "US-ONLY": 3}
    jobs.sort(key=lambda j: order.get(j.get("region"), 9))

    # Применяем к пулу ТЕКУЩИЕ фильтры из job_finder (роль + стоп-слова + сеньорность),
    # чтобы уже собранные нерелевантные вакансии не доедались впустую.
    _loose = os.environ.get("LOOSE_FILTER") == "1"
    try:
        import job_finder as _jf
        def _passes(j):
            t, d = j.get("title", ""), j.get("description", "")
            return ((not _jf.blocked(t)) and _jf.remote_ok(t, d, None)
                    and (_loose or _jf.matches_role(t)))
    except Exception:
        def _passes(j):
            return True

    def _fresh(j):
        return (j.get("region") in PROCESS_REGIONS
                and j.get("url") not in processed
                and _passes(j))

    # Свежие — первыми: где новее вакансия, там меньше откликов/конкурентов.
    _today = datetime.date.today().isoformat()
    queue = sorted((j for j in jobs if _fresh(j)),
                   key=lambda j: j.get("date") or _today, reverse=True)[:MAX_JOBS]
    print(f"К обработке: {len(queue)} (необработанных в пуле всего: "
          f"{sum(1 for j in jobs if _fresh(j))})")

    done = 0
    for job in queue:
        title = job.get("title", "")[:50]
        try:
            # Шаг 1 (дёшево, Haiku): предварительная оценка fit
            pre = score_fit(job)
            pre_score = int(pre.get("fit_score", 0) or 0)
            if pre_score < PRE_MIN_FIT:
                mark_processed(job.get("url", ""))
                print(f"  [{pre_score:>3}] {title} — предфильтр отсеял (Haiku)")
                time.sleep(1)
                continue

            # Шаг 2 (дорого, Sonnet): полный тюнинг только для прошедших
            res = analyze(job, master_txt)
            score = int(res.get("fit_score", 0) or 0)
            if score < MIN_FIT:
                mark_processed(job.get("url", ""))
                print(f"  [{score:>3}] {title} — ниже порога после тюнинга, пропуск")
                time.sleep(1)
                continue
            folder = write_package(job, res, master_txt)
            log_application(job, res, folder)
            mark_processed(job.get("url", ""))
            done += 1
            print(f"  [{score:>3}] {title} -> {folder}/")
            time.sleep(1)
        except Exception as e:
            print(f"  [ERR] {title}: {e}")

    print(f"\nГотово: {done} комплектов в ./{REVIEW_DIR}/")
    print("В каждой папке: подогнанное резюме (.docx) + cover_letter.txt + ats_report.md")


if __name__ == "__main__":
    main()
